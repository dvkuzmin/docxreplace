import os
import win32api, win32con
import docx
import re

def file_is_hidden(p):
    if os.name=='nt':
        attribute = win32api.GetFileAttributes(p)
        return attribute & (win32con.FILE_ATTRIBUTE_HIDDEN | win32con.FILE_ATTRIBUTE_SYSTEM)


def searching_files(path, pattern, new_text):
	file_list = [f for f in os.listdir(path) if not file_is_hidden(os.path.join(path, f))]
	for i in file_list:
		if os.path.isdir(os.path.join(path, i)):
			searching_files(os.path.join(path, i), pattern, new_text)
		elif i.endswith('.docx'):
			replace_pattern(os.path.join(path, i), pattern, new_text)


def replace_pattern(file, pattern=r'', new_text=''):
	doc = docx.Document(file)
	for paragraph in doc.paragraphs:
		paragraph.text = re.sub(pattern, new_text, paragraph.text)
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				cell.text = re.sub(pattern, new_text, cell.text)
	doc.save(file)